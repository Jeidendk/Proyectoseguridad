
import docx
from docx.shared import Pt, RGBColor
import os
import re

# Constantes
TEMPLATE_FILE = 'plantilla-titulos-normas-apa-7-mentur (15).docx'
MARKDOWN_FILE = 'INFORME_ACADEMICO.md'
OUTPUT_FILE = 'INFORME_ACADEMICO_FINAL_v3.docx'

def safe_add_paragraph(doc, text, style_name):
    try:
        doc.add_paragraph(text, style=style_name)
    except KeyError:
        # Fallback handling
        try:
            p = doc.add_paragraph(text)
            if style_name == 'List Bullet':
                p.style = 'Normal'
                p.text = "• " + text
            elif style_name == 'List Number':
                p.style = 'Normal'
        except Exception as e:
            print(f"Error al agregar párrafo fallback: {e}")

def add_code_block(doc, code_text):
    """Adds a paragraph formatted as a code block."""
    try:
        # Create a new paragraph
        p = doc.add_paragraph()
        
        # Add run with code text
        run = p.add_run(code_text)
        
        # Apply formatting manually since we might not have a 'Code' style
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30) # Dark gray
        
        # Try to set paragraph properties if possible (requires OXML usually for shading)
        # For simplicity, we stick to font styling to denote code
        
    except Exception as e:
        print(f"Error agregando bloque de código: {e}")

def parse_markdown_to_docx(md_path, template_path, output_path):
    if not os.path.exists(template_path):
        print(f"Error: No se encuentra la plantilla {template_path}")
        return
    
    if not os.path.exists(md_path):
        print(f"Error: No se encuentra el archivo Markdown {md_path}")
        return

    try:
        doc = docx.Document(template_path)
        print(f"Plantilla cargada.") 
    except Exception as e:
        print(f"Error al cargar la plantilla: {e}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    print("Procesando contenido...")
    
    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()
        
        # Code block detection (triple backticks)
        if stripped.startswith("```"):
            if in_code_block:
                # End of block
                in_code_block = False
                if code_buffer:
                    full_code = "".join(code_buffer).rstrip()
                    add_code_block(doc, full_code)
                    code_buffer = []
            else:
                # Start of block
                in_code_block = True
                code_buffer = []
            continue
            
        if in_code_block:
            code_buffer.append(line)
            continue

        # Normal text processing
        if not stripped:
            continue
            
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            safe_add_paragraph(doc, stripped[2:], 'List Bullet')
        elif stripped.startswith('1. '):
            content = re.sub(r'^\d+\.\s+', '', stripped)
            safe_add_paragraph(doc, content, 'List Number')
        elif stripped.startswith('<!--'):
            pass
        elif stripped.startswith('---'):
            doc.add_page_break()
        elif stripped.startswith('>'):
             try:
                p = doc.add_paragraph(stripped[1:].strip())
                p.italic = True
             except:
                doc.add_paragraph(stripped[1:].strip())
        else:
            # Basic table handling (detected by pipe)
            if '|' in stripped and stripped.count('|') > 1:
                # Naive table skipped for now/rendered as text to avoid complexities
                # Or better: render as monospaced to align
                if '---' in stripped: # separator line
                    continue 
                add_code_block(doc, stripped)
            else:
                doc.add_paragraph(stripped)

    try:
        doc.save(output_path)
        print(f"Documento generado exitosamente: {output_path}")
    except Exception as e:
        print(f"Error al guardar el documento: {e}")

if __name__ == "__main__":
    parse_markdown_to_docx(MARKDOWN_FILE, TEMPLATE_FILE, OUTPUT_FILE)
